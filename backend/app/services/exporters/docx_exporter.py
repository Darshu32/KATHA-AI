"""DOCX specification document — python-docx layout.

Produces a Word file with the same sections as the PDF but editable and
client-deliverable via email.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from app.services.exporters._synthesis import (
    derive_assembly_instructions,
    derive_maintenance_guide,
)


def _fmt_range(v) -> str:
    if isinstance(v, (list, tuple)) and len(v) == 2:
        return f"{v[0]} – {v[1]}"
    if v is None:
        return "—"
    return str(v)


def _flat(v) -> str:
    if isinstance(v, dict):
        return ", ".join(f"{k}: {_flat(x)}" for k, x in v.items())
    if isinstance(v, list):
        return "; ".join(_flat(x) for x in v)
    if isinstance(v, tuple):
        return _fmt_range(list(v))
    return "—" if v is None else str(v)


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x1D, 0x1A)


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x4A, 0x46, 0x3F)


def _table(doc, header: list[str], rows: list[list]):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v) if v is not None else "—"
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    return table


def export(spec: dict, graph: dict) -> dict:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    meta = spec["meta"]

    title = doc.add_paragraph()
    run = title.add_run("KATHA Design Dossier")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"{meta['project_name']} · {meta.get('room_type', '—')} · Theme: {meta.get('theme', '—')} · "
        f"{meta['dimensions_m'].get('length','?')} x {meta['dimensions_m'].get('width','?')} x "
        f"{meta['dimensions_m'].get('height','?')} m · Generated {meta['generated_at']}"
    ).italic = True

    # Materials.
    _h1(doc, "Materials")
    mat = spec["material"]
    header = ["Name", "Grade", "Finish", "Color", "Supplier", "Lead (wk)", "Cost INR"]
    for label, rows_src in [
        ("Primary Structure", mat["primary_structure"]),
        ("Secondary", mat["secondary_materials"]),
        ("Upholstery", mat["upholstery"]),
        ("Hardware", mat["hardware"]),
        ("Finishing", mat["finishing"]),
    ]:
        if not rows_src:
            continue
        _h2(doc, label)
        rows = [
            [
                r.get("name", "—"),
                r.get("grade", "—"),
                r.get("finish", "—"),
                r.get("color", "—"),
                r.get("supplier", "—"),
                _fmt_range(r.get("lead_time_weeks")),
                _fmt_range(r.get("cost_inr")),
            ]
            for r in rows_src
        ]
        _table(doc, header, rows)
    tn = mat.get("total_notes", {})
    if tn.get("adjusted_note"):
        p = doc.add_paragraph(tn["adjusted_note"])
        p.runs[0].font.size = Pt(9)

    # Manufacturing.
    doc.add_page_break()
    _h1(doc, "Manufacturing")
    for trade, block in spec["manufacturing"].items():
        _h2(doc, trade.replace("_", " ").title())
        _table(doc, ["Key", "Value"], [[k.replace("_", " ").title(), _flat(v)] for k, v in block.items()])

    # MEP.
    doc.add_page_break()
    _h1(doc, "MEP — Mechanical, Electrical, Plumbing")
    for system, block in spec["mep"].items():
        _h2(doc, system.upper())
        _table(doc, ["Parameter", "Value"], [[k.replace("_", " ").title(), _flat(v)] for k, v in block.items()])

    # Cost.
    doc.add_page_break()
    _h1(doc, "Cost Estimate")
    c = spec["cost"]
    totals = c.get("totals", {})
    doc.add_paragraph(
        f"Status: {c.get('status','pending').title()} · Currency: {c.get('currency','INR')} · "
        f"Total (low / base / high): {totals.get('low','—')} / {totals.get('base','—')} / {totals.get('high','—')}"
    )
    line_items = c.get("line_items") or []
    if line_items:
        rows = [
            [
                li.get("category", "—"),
                li.get("itemName") or li.get("item_name") or li.get("name", "—"),
                str(li.get("quantity", "—")),
                li.get("unit", "—"),
                _fmt_range(li.get("unitRate") or li.get("unit_rate")),
                _fmt_range(li.get("totalLow") or li.get("total_low")),
                _fmt_range(li.get("totalHigh") or li.get("total_high")),
            ]
            for li in line_items[:60]
        ]
        _table(doc, ["Category", "Item", "Qty", "Unit", "Rate", "Low", "High"], rows)
    else:
        doc.add_paragraph("Cost will be computed after the estimation pipeline completes.")
    for a in c.get("assumptions", [])[:10]:
        doc.add_paragraph(f"• {a}")

    # Stage 14 — BRD §5A explicitly asks for assembly instructions and a
    # maintenance & care guide in the DOCX deliverable. Both sections
    # synthesize from spec content via app.services.exporters._synthesis.
    _build_assembly_section(doc, spec, graph)
    _build_maintenance_section(doc, spec, graph)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return {
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": f"{_safe_name(meta['project_name'])}-dossier.docx",
        "bytes": buffer.getvalue(),
    }


def _safe_name(name: str) -> str:
    return "".join(c if c.isalnum() or c in "-_" else "-" for c in name).strip("-") or "project"


# ─────────────────────────────────────────────────────────────────────
# Stage 14 — Assembly Instructions + Maintenance & Care Guide
# ─────────────────────────────────────────────────────────────────────


def _build_assembly_section(doc, spec: dict, graph: dict) -> None:
    guide = derive_assembly_instructions(spec, graph)
    doc.add_page_break()
    _h1(doc, "Assembly Instructions")

    summary = doc.add_paragraph(guide["summary"])
    summary.runs[0].font.size = Pt(10)

    if guide["tolerance_notes"]:
        _h2(doc, "Tolerances at a glance")
        for note in guide["tolerance_notes"]:
            p = doc.add_paragraph(f"• {note}")
            p.runs[0].font.size = Pt(9)

    _h2(doc, "Step-by-step")
    if not guide["steps"]:
        doc.add_paragraph(
            "Sequence not yet generated; run the manufacturing spec to populate."
        )
    else:
        for step in guide["steps"]:
            p = doc.add_paragraph()
            run = p.add_run(f"Step {step['step_number']}. ")
            run.bold = True
            run.font.size = Pt(11)
            p.add_run(step["action"]).font.size = Pt(11)

            tools_p = doc.add_paragraph(f"Tools: {', '.join(step['tools'])}")
            tools_p.runs[0].font.size = Pt(9)
            tools_p.runs[0].italic = True

            safety_p = doc.add_paragraph(f"Note: {step['safety']}")
            safety_p.runs[0].font.size = Pt(9)

    if guide["qa_gates"]:
        _h2(doc, "QA gates (sign-off required)")
        _table(
            doc,
            ["Gate", "What to verify"],
            [[g.get("gate", "—"), g.get("description", "—") if isinstance(g, dict) else str(g)]
             for g in guide["qa_gates"]],
        )

    _h2(doc, "Packaging")
    p = doc.add_paragraph(guide["packaging"])
    p.runs[0].font.size = Pt(10)


def _build_maintenance_section(doc, spec: dict, graph: dict) -> None:
    guide = derive_maintenance_guide(spec, graph)
    doc.add_page_break()
    _h1(doc, "Maintenance & Care Guide")

    intro = doc.add_paragraph(guide["intro"])
    intro.runs[0].font.size = Pt(10)

    if not guide["categories"]:
        doc.add_paragraph(
            "No materials matched the care matrix — contact the studio for "
            "a customised care plan."
        )
    else:
        for entry in guide["categories"]:
            _h2(
                doc,
                f"{entry['category'].title()}  —  {', '.join(entry['applies_to'])}",
            )
            for label, items in (
                ("Daily", entry["daily"]),
                ("Weekly", entry["weekly"]),
                ("Monthly", entry["monthly"]),
                ("Annually", entry["annually"]),
            ):
                if not items:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run("; ".join(items)).font.size = Pt(10)
            if entry["warnings"]:
                p = doc.add_paragraph()
                run = p.add_run("Warnings: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
                run.font.size = Pt(10)
                p.add_run("; ".join(entry["warnings"])).font.size = Pt(10)

    _h2(doc, "General notes")
    for note in guide["general_notes"]:
        p = doc.add_paragraph(f"• {note}")
        p.runs[0].font.size = Pt(9)
