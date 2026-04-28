"""DOCX importer — paragraphs, headings, tables (python-docx)."""

from __future__ import annotations

import io
import re
from typing import Any

try:
    from docx import Document
except Exception:  # noqa: BLE001
    Document = None

from app.services.importers.text_importer import _detect_brief_signals


def parse(filename: str, payload: bytes) -> dict[str, Any]:
    if Document is None:
        return {
            "format": "docx",
            "filename": filename,
            "size_bytes": len(payload),
            "summary": "python-docx not installed.",
            "extracted": {},
            "warnings": ["python-docx is not available — install to parse .docx files."],
        }
    try:
        doc = Document(io.BytesIO(payload))
    except Exception as exc:  # noqa: BLE001
        return {
            "format": "docx",
            "filename": filename,
            "size_bytes": len(payload),
            "summary": "Could not parse DOCX.",
            "extracted": {},
            "warnings": [f"python-docx load failed: {exc}"],
        }

    paragraphs: list[dict[str, Any]] = []
    headings: list[str] = []
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.startswith("Heading"):
            headings.append(txt)
        paragraphs.append({"style": style, "text": txt})

    tables: list[dict[str, Any]] = []
    for tbl in doc.tables[:10]:
        rows = []
        for row in tbl.rows[:25]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({
            "row_count": len(tbl.rows),
            "col_count": len(tbl.columns),
            "sample": rows,
        })

    full_text = "\n".join(p["text"] for p in paragraphs)
    signals = _detect_brief_signals(full_text)

    return {
        "format": "docx",
        "filename": filename,
        "size_bytes": len(payload),
        "summary": (
            f"DOCX: {len(paragraphs)} paragraph(s); {len(headings)} heading(s); "
            f"{len(tables)} table(s)."
        ),
        "extracted": {
            "headings": headings,
            "paragraph_count": len(paragraphs),
            "text_excerpt": full_text[:3000],
            "tables": tables,
            "brief_signals": signals,
        },
        "warnings": [],
    }
