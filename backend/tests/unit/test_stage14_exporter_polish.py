"""Stage 14 unit tests — exporter polish closing the BRD §5A gaps.

Three gaps closed in this stage; each gets focused coverage here:

1. **DOCX assembly + maintenance** — synthesis derivers produce the
   right structure; the DOCX exporter writes the new sections (we
   verify by reopening the file via python-docx and asserting the
   heading text appears).
2. **PPTX assembly + maintenance + render embed** — synthesis derivers
   work; PPTX zip contains slides 7 + 8 (assembly + maintenance) and
   embedded media when renders are passed in.
3. **HTML inline GLTF** — the model-viewer ``src`` attribute carries a
   base64 ``data:model/gltf+json`` URI when graph geometry is present;
   the previous "drop the GLTF next to the file" warning is gone.

All tests run in-process — no DB, no network, no Celery.
"""

from __future__ import annotations

import base64
import io
import zipfile

import pytest

from app.services.exporters._synthesis import (
    collect_render_images,
    derive_assembly_instructions,
    derive_maintenance_guide,
)


# ─────────────────────────────────────────────────────────────────────
# Fixture builders
# ─────────────────────────────────────────────────────────────────────


def _spec_with_full_manufacturing() -> dict:
    return {
        "meta": {
            "project_name": "Test Suite Project",
            "generated_at": "2026-05-02T00:00:00+00:00",
            "room_type": "study",
            "theme": "mid_century",
            "dimensions_m": {"length": 4.0, "width": 3.0, "height": 3.0},
        },
        "material": {
            "primary_structure": [
                {"name": "Walnut", "category": "wood", "grade": "Seasoned grade A",
                 "finish": "Natural oil", "color": "warm brown",
                 "supplier": "Local fabricator", "lead_time_weeks": [2, 4],
                 "cost_inr": [300, 800]},
            ],
            "secondary_materials": [
                {"name": "Mild Steel", "category": "metal",
                 "finish": "Powder coat", "lead_time_weeks": [3, 6],
                 "cost_inr": [60, 90]},
            ],
            "hardware": [
                {"name": "Brass Knob", "supplier": "—", "lead_time_weeks": None,
                 "cost_inr": [500, 2000]},
            ],
            "upholstery": [
                {"name": "Leather Grade A", "category": "leather",
                 "lead_time_weeks": [4, 6], "cost_inr": [800, 3000]},
            ],
            "finishing": [
                {"name": "Lacquer Topcoat", "category": "finish",
                 "lead_time_weeks": [1, 2], "cost_inr": [50, 150]},
            ],
            "total_notes": {"waste_factor_pct": 12, "adjusted_note": "Apply 10-15% waste."},
        },
        "manufacturing": {
            "woodworking": {
                "tolerance_structural_mm": 1.0,
                "tolerance_cosmetic_mm": 2.0,
                "joinery_recommended": [{"method": "mortise_tenon"}],
                "lead_time_weeks": [4, 8],
            },
            "metal_fabrication": {
                "tolerance_structural_mm": 1.0,
                "tolerance_cosmetic_mm": 2.0,
                "lead_time_weeks": [6, 10],
            },
            "upholstery": {
                "webbing_tension_kg_per_inch": [5, 8],
                "stitch_density_per_inch": [4, 6],
                "foam_tolerance_mm": 5,
                "lead_time_weeks": [3, 6],
            },
            "assembly": {
                "sequence": [
                    "Frame + joinery dry-fit",
                    "Sand + finish while separable",
                    "Hardware install (torque per spec)",
                    "Upholstery mount",
                    "Final QC + packaging",
                ],
                "qa_gates": [
                    {"gate": "Stability", "description": "No wobble at 80 kg load"},
                    {"gate": "Surface", "description": "No visible glue beads"},
                ],
                "packaging": "Edge-protected corrugate; corner foam.",
            },
        },
        "mep": {
            "hvac": {"room_volume_m3": 36.0, "cfm_required": 240, "cost_inr": 45000},
            "electrical": {"lux_target": 300, "cost_inr": 18000},
            "plumbing": {"fixture_count": 2, "cost_inr": 12000},
        },
        "cost": {
            "status": "ready",
            "currency": "INR",
            "totals": {"low": 50000, "base": 65000, "high": 80000},
            "line_items": [
                {"category": "Material", "item_name": "Walnut top",
                 "quantity": 1, "unit": "piece", "unit_rate": [300, 800],
                 "total_low": 30000, "total_high": 50000},
            ],
            "assumptions": ["12% waste factor applied."],
        },
        "objects_count": 2,
    }


def _graph_with_geometry() -> dict:
    return {
        "room": {"type": "study",
                 "dimensions": {"length": 4.0, "width": 3.0, "height": 3.0}},
        "objects": [
            {"id": "desk", "type": "table", "name": "Walnut Desk",
             "dimensions": {"length": 1.4, "width": 0.6, "height": 0.75},
             "position": {"x": 1.0, "y": 0, "z": 1.0},
             "rotation": {"x": 0, "y": 0, "z": 0}, "color": "#5C3B1E"},
            {"id": "chair", "type": "chair", "name": "Studio Chair",
             "dimensions": {"length": 0.5, "width": 0.5, "height": 0.85},
             "position": {"x": 1.5, "y": 0, "z": 1.6},
             "rotation": {"x": 0, "y": 0, "z": 0}, "color": "#3D3A36"},
        ],
        "style": {"signature_moves": ["Tapered legs", "Brass accents", "Low profile"]},
        "materials": [
            {"name": "Walnut", "category": "wood"},
            {"name": "Mild Steel", "category": "metal"},
        ],
    }


_PNG_1X1 = base64.b64decode(
    # 1x1 transparent PNG — the smallest valid image we can ship.
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


# ─────────────────────────────────────────────────────────────────────
# _synthesis: assembly instructions
# ─────────────────────────────────────────────────────────────────────


def test_assembly_synthesis_returns_step_count_matching_sequence():
    spec = _spec_with_full_manufacturing()
    guide = derive_assembly_instructions(spec, {})
    assert len(guide["steps"]) == 5
    assert guide["steps"][0]["step_number"] == 1
    assert guide["steps"][-1]["step_number"] == 5


def test_assembly_synthesis_attaches_tools_and_safety_per_step():
    spec = _spec_with_full_manufacturing()
    guide = derive_assembly_instructions(spec, {})
    for step in guide["steps"]:
        assert step["tools"], "every step should have at least one tool"
        assert step["safety"], "every step should carry a safety note"


def test_assembly_synthesis_includes_qa_gates_and_packaging():
    spec = _spec_with_full_manufacturing()
    guide = derive_assembly_instructions(spec, {})
    assert len(guide["qa_gates"]) == 2
    assert "corrugate" in guide["packaging"].lower()


def test_assembly_synthesis_tolerance_notes_pulled_from_manufacturing():
    spec = _spec_with_full_manufacturing()
    guide = derive_assembly_instructions(spec, {})
    text = "\n".join(guide["tolerance_notes"]).lower()
    assert "woodworking" in text
    assert "metal" in text
    assert "upholstery" in text


def test_assembly_synthesis_handles_missing_manufacturing_block():
    guide = derive_assembly_instructions({}, {})
    assert guide["steps"] == []
    assert "generated once the manufacturing spec" in guide["summary"].lower()


# ─────────────────────────────────────────────────────────────────────
# _synthesis: maintenance guide
# ─────────────────────────────────────────────────────────────────────


def test_maintenance_guide_includes_only_present_categories():
    spec = _spec_with_full_manufacturing()
    guide = derive_maintenance_guide(spec, {})
    cats = {c["category"] for c in guide["categories"]}
    # Spec has wood, metal, leather, finishing, hardware → all should appear.
    assert "wood" in cats
    assert "metal" in cats
    assert "leather" in cats
    assert "finish" in cats
    assert "hardware" in cats


def test_maintenance_guide_warnings_not_empty_for_known_categories():
    spec = _spec_with_full_manufacturing()
    guide = derive_maintenance_guide(spec, {})
    leather = next(c for c in guide["categories"] if c["category"] == "leather")
    assert leather["warnings"]
    assert any("sun" in w.lower() or "uv" in w.lower() or "alcohol" in w.lower()
               for w in leather["warnings"])


def test_maintenance_guide_general_notes_always_present():
    spec = _spec_with_full_manufacturing()
    guide = derive_maintenance_guide(spec, {})
    assert len(guide["general_notes"]) >= 2


def test_maintenance_guide_handles_empty_material_block():
    guide = derive_maintenance_guide({}, {})
    assert guide["categories"] == []
    assert "pending material spec" in guide["intro"].lower() or \
           "generated once" in guide["intro"].lower()


def test_maintenance_guide_classifies_by_name_when_category_missing():
    spec = {
        "material": {
            "primary_structure": [{"name": "Walnut Slab"}],  # no category
            "secondary_materials": [],
            "hardware": [],
            "upholstery": [],
            "finishing": [],
        }
    }
    guide = derive_maintenance_guide(spec, {})
    assert any(c["category"] == "wood" for c in guide["categories"])


# ─────────────────────────────────────────────────────────────────────
# Render collection
# ─────────────────────────────────────────────────────────────────────


def test_collect_renders_returns_empty_when_no_renders():
    assert collect_render_images({}, {}) == []
    assert collect_render_images({"renders": []}, {"renders": []}) == []


def test_collect_renders_picks_up_spec_renders():
    spec = {"renders": [
        {"caption": "Hero", "mime": "image/png", "bytes": _PNG_1X1},
    ]}
    rendered = collect_render_images(spec, {})
    assert len(rendered) == 1
    assert rendered[0]["caption"] == "Hero"
    assert rendered[0]["ext"] == ".png"


def test_collect_renders_picks_up_graph_assets_with_render_kind():
    graph = {"assets": [
        {"kind": "render_2d", "caption": "North view",
         "mime": "image/png", "bytes": _PNG_1X1},
        {"kind": "scene_3d", "bytes": b"\x00\x00"},  # ignored
    ]}
    rendered = collect_render_images({}, graph)
    assert len(rendered) == 1
    assert rendered[0]["caption"] == "North view"


def test_collect_renders_drops_unknown_mime_silently():
    graph = {"renders": [
        {"caption": "weird", "mime": "image/tiff", "bytes": b"x"},
    ]}
    assert collect_render_images({}, graph) == []


# ─────────────────────────────────────────────────────────────────────
# DOCX exporter — assembly + maintenance sections appear
# ─────────────────────────────────────────────────────────────────────


def test_docx_export_contains_assembly_and_maintenance_headings():
    pytest.importorskip("docx")  # python-docx
    from docx import Document
    from app.services.exporters import docx_exporter

    spec = _spec_with_full_manufacturing()
    graph = _graph_with_geometry()
    out = docx_exporter.export(spec, graph)

    assert out["filename"].endswith(".docx")
    assert out["bytes"]

    # Reopen the docx and verify the new headings appear.
    doc = Document(io.BytesIO(out["bytes"]))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    joined = " | ".join(headings)
    assert "Assembly Instructions" in joined
    assert "Maintenance & Care Guide" in joined
    # Check the step content actually made it in.
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Step 1" in body
    assert "Frame + joinery dry-fit" in body


def test_docx_export_handles_missing_manufacturing_gracefully():
    pytest.importorskip("docx")
    from docx import Document
    from app.services.exporters import docx_exporter

    spec = _spec_with_full_manufacturing()
    spec["manufacturing"] = {"assembly": {}}  # strip the sequence
    out = docx_exporter.export(spec, {})

    doc = Document(io.BytesIO(out["bytes"]))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Assembly Instructions" in body  # heading still rendered
    assert "Sequence not yet generated" in body


# ─────────────────────────────────────────────────────────────────────
# PPTX exporter — slide count, assembly + maintenance, embedded image
# ─────────────────────────────────────────────────────────────────────


def _open_pptx(zip_bytes: bytes) -> zipfile.ZipFile:
    return zipfile.ZipFile(io.BytesIO(zip_bytes), "r")


def test_pptx_export_emits_nine_slides_with_assembly_and_maintenance():
    from app.services.exporters import pptx_exporter

    spec = _spec_with_full_manufacturing()
    out = pptx_exporter.export(spec, _graph_with_geometry())
    zf = _open_pptx(out["bytes"])
    slide_files = sorted(n for n in zf.namelist() if n.startswith("ppt/slides/slide"))
    # Stage 14 added Assembly + Maintenance — total = 9 slides.
    assert len(slide_files) == 9, slide_files

    # Slide 7 = Assembly, Slide 8 = Maintenance (cover, concept, spec,
    # cost, timeline, drawings, assembly, maintenance, next-steps).
    s7 = zf.read("ppt/slides/slide7.xml").decode()
    assert "Assembly instructions" in s7
    assert "Frame + joinery dry-fit" in s7

    s8 = zf.read("ppt/slides/slide8.xml").decode()
    assert "Maintenance &amp; care" in s8 or "Maintenance & care" in s8


def test_pptx_export_embeds_render_image_when_provided():
    from app.services.exporters import pptx_exporter

    spec = _spec_with_full_manufacturing()
    spec["renders"] = [
        {"caption": "Hero render", "mime": "image/png", "bytes": _PNG_1X1},
        {"caption": "Plan", "mime": "image/png", "bytes": _PNG_1X1},
    ]
    out = pptx_exporter.export(spec, _graph_with_geometry())
    zf = _open_pptx(out["bytes"])
    media = [n for n in zf.namelist() if n.startswith("ppt/media/")]
    assert len(media) == 2
    # Concept slide (slide 2) should reference the hero render.
    s2_rels = zf.read("ppt/slides/_rels/slide2.xml.rels").decode()
    assert "render_1.png" in s2_rels
    # Drawings index slide (slide 6) should reference the gallery.
    s6_rels = zf.read("ppt/slides/_rels/slide6.xml.rels").decode()
    assert "render_1.png" in s6_rels or "render_2.png" in s6_rels


def test_pptx_export_works_without_renders_no_media_part():
    from app.services.exporters import pptx_exporter

    out = pptx_exporter.export(_spec_with_full_manufacturing(), {})
    zf = _open_pptx(out["bytes"])
    media = [n for n in zf.namelist() if n.startswith("ppt/media/")]
    assert media == []
    # Slide rels for slide 2 should NOT reference any image.
    s2_rels = zf.read("ppt/slides/_rels/slide2.xml.rels").decode()
    assert "image" not in s2_rels.lower()


def test_pptx_content_types_registers_image_extensions_only_when_used():
    from app.services.exporters import pptx_exporter

    no_renders = pptx_exporter.export(_spec_with_full_manufacturing(), {})
    ct = zipfile.ZipFile(io.BytesIO(no_renders["bytes"])).read("[Content_Types].xml").decode()
    assert 'Extension="png"' not in ct

    spec = _spec_with_full_manufacturing()
    spec["renders"] = [{"caption": "x", "mime": "image/png", "bytes": _PNG_1X1}]
    with_renders = pptx_exporter.export(spec, {})
    ct2 = zipfile.ZipFile(io.BytesIO(with_renders["bytes"])).read("[Content_Types].xml").decode()
    assert 'Extension="png"' in ct2


# ─────────────────────────────────────────────────────────────────────
# HTML exporter — inline GLTF data URI
# ─────────────────────────────────────────────────────────────────────


def test_html_export_embeds_gltf_data_uri_when_geometry_present():
    from app.services.exporters import html_exporter

    spec = _spec_with_full_manufacturing()
    graph = _graph_with_geometry()
    out = html_exporter.export(spec, graph)

    body = out["bytes"].decode("utf-8")
    assert 'src="data:model/gltf+json;base64,' in body
    # The legacy "drop the GLTF next to this HTML" warning must be gone.
    assert "Drop the GLTF" not in body


def test_html_export_falls_back_when_no_geometry():
    from app.services.exporters import html_exporter

    out = html_exporter.export(_spec_with_full_manufacturing(), {})
    body = out["bytes"].decode("utf-8")
    # Empty src + a graceful pending-message instead of the GLTF.
    assert 'src=""' in body
    assert "preview will appear" in body or "preview pending" in body


def test_html_export_filename_remains_stable():
    from app.services.exporters import html_exporter

    out = html_exporter.export(_spec_with_full_manufacturing(), _graph_with_geometry())
    assert out["filename"].endswith("-viewer.html")
    assert out["content_type"].startswith("text/html")
